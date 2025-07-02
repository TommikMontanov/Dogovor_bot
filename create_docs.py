from io import BytesIO
import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from tempfile import NamedTemporaryFile

def set_font(cell, text):
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

def num_to_words(number):
    try:
        from num2words import num2words
        return num2words(number, lang='ru').capitalize()
    except:
        return f"{number:,.2f}".replace(",", " ").replace(".", ",")


def generate_excel_and_docx(data):
    # Сбор таблицы
    rows = []
    for i, item in enumerate(data, 2):
        rows.append({
            "№": i - 1,
            "Наименование": f"Дорожный знак,{item['id']},размер {item['размер']}",
            "Ед. изм": "шт",
            "Кол-во": item["кол-во"],
            "Цена": item["цена"],
            "Стоимость поставки": "",
            "НДС ставка": "12%",
            "НДС сумма": "",
            "Всего сумма с НДС": ""
        })

    df = pd.DataFrame(rows)

    last_row_excel = len(df) + 2
    df.loc[len(df.index)] = {
        "№": "ИТОГО:",
        "Наименование": "",
        "Ед. изм": "",
        "Кол-во": f"=SUM(D2:D{last_row_excel - 1})",
        "Цена": "",
        "Стоимость поставки": f"=SUM(F2:F{last_row_excel - 1})",
        "НДС ставка": "",
        "НДС сумма": f"=SUM(H2:H{last_row_excel - 1})",
        "Всего сумма с НДС": f"=SUM(I2:I{last_row_excel - 1})"
    }

    with NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        df.to_excel(tmp.name, index=False)
        wb = load_workbook(tmp.name)
        ws = wb.active

        # Формулы
        for i in range(2, last_row_excel):
            ws[f"F{i}"] = f"=D{i}*E{i}"
            ws[f"H{i}"] = f"=F{i}*0.12"
            ws[f"I{i}"] = f"=F{i}+H{i}"

        # Стили
        thin = Side(style="thin")
        thick = Side(style="thick")
        thin_border = Border(left=thin, right=thin, top=thin, bottom=thin)
        thick_border = Border(left=thick, right=thick, top=thick, bottom=thick)

        for row in ws.iter_rows():
            for cell in row:
                row_idx = cell.row
                cell.alignment = Alignment(wrap_text=True, horizontal="center", vertical="center")
                cell.font = Font(name="Arial", size=10, bold=(row_idx == 1 or row_idx == last_row_excel))
                cell.border = thick_border if row_idx in [1, last_row_excel] else thin_border

        for col in ws.columns:
            max_len = max(len(str(cell.value)) if cell.value else 0 for cell in col)
            ws.column_dimensions[col[0].column_letter].width = max_len + 2

        for row in ws.iter_rows():
            for cell in row:
                if cell.column_letter == "B":
                    ws.row_dimensions[cell.row].height = 30

        excel_buffer = BytesIO()
        wb.save(excel_buffer)
        excel_buffer.seek(0)

    # DOCX генерация
    total_qty = sum(item["кол-во"] for item in data)
    total_cost = sum(item["кол-во"] * item["цена"] for item in data)
    nds = int(total_cost * 0.12)
    total_with_nds = total_cost + nds

    today = datetime.now()
    today_str = today.strftime("%d/%m")
    date_for_contract = f"«{today.day}» {today.strftime('%B')} {today.year}г."

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(9)

    doc.add_paragraph(f"Договор-счет №   {today_str}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"на выполнение работ и услуг").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("г. Ташкент\t\t\t\t\t\t\t\t\t" + date_for_contract)

    doc.add_paragraph(
        '\t\tOOО «NUR  FAYZ  REKLAMA», именуемый в дальнейшем «Исполнитель», в лице директора Файзиева Н.С., '
        'действующий на основании Устава, с одной стороны, и "HIGH CITY-DEVELOPERS" MCHJ, именуемый в дальнейшем '
        '«Заказчик», в лице директора\t\t\t, действующего на основании Устава, с другой стороны, заключили настоящий Договор о нижеследующем:'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('1. ПРЕДМЕТ ДОГОВОРА')
    run.bold = True

    doc.add_paragraph(
        "1.1. Исполнитель обязуется изготовить продукцию и прочие услуги, а Заказчик обязуется оплатить и принять на условиях, "
        "установленных настоящим Договором согласно следующей спецификации:"
    )
# Создаем таблицу 2 колонки, 1 строка
    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    headers = ["№", "Наименование", "Ед. изм", "Кол-во", "Цена", "Стоимость поставки", "НДС ставка", "НДС сумма", "Всего сумма с НДС"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for i, item in enumerate(data, 1):
        qty = item["кол-во"]
        price = item["цена"]
        cost = qty * price
        vat = int(cost * 0.12)
        total = cost + vat
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = f"Дорожный знак {item['id']}, размер {item['размер']}"
        row_cells[2].text = "шт"
        row_cells[3].text = str(qty)
        row_cells[4].text = f"{price:,.2f}".replace(",", " ").replace(".", ",")
        row_cells[5].text = f"{cost:,.2f}".replace(",", " ").replace(".", ",")
        row_cells[6].text = "12%"
        row_cells[7].text = f"{vat:,.2f}".replace(",", " ").replace(".", ",")
        row_cells[8].text = f"{total:,.2f}".replace(",", " ").replace(".", ",")

    row_cells = table.add_row().cells
    row_cells[0].text = "ИТОГО:"
    row_cells[3].text = str(total_qty)
    row_cells[5].text = f"{total_cost:,.2f}".replace(",", " ").replace(".", ",")
    row_cells[7].text = f"{nds:,.2f}".replace(",", " ").replace(".", ",")
    row_cells[8].text = f"{total_with_nds:,.2f}".replace(",", " ").replace(".", ",")

    doc.add_paragraph(
        'Данная спецификация одновременно является протоколом согласования цены при данном тираже. '
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"2. ОБЩАЯ СУММА И ПОРЯДОК ОПЛАТЫ\n")
    run.bold = True

    doc.add_paragraph(
        f"2.1. Общая сумма настоящего Договора составляет: {total_with_nds:,.2f} "
        f"({num_to_words(total_with_nds)} сумов 00 тийин), с учетом НДС."
    )
    # Пункт 2.2
    doc.add_paragraph(
        '2.2. Началом выполнения работ является предоплата в размере 100 % от суммы, указанной в договоре, '
        'путем перечисления «Заказчиком» на расчетный счет «Исполнителя» в течение 10 (десяти) банковских дней с момента регистрации настоящего договора.'
    )

    # Заголовок раздела 3
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('3. УСЛОВИЯ ИЗГОТОВЛЕНИЯ И ПОСТАВКИ ПРОДУКЦИИ')
    run.bold = True

    doc.add_paragraph(
        '3.1. Срок изготовления и поставки продукции — 10 рабочих дней со дня поступления предоплаты на расчетный счет '
        'Исполнителя и утверждения оригинал-макета со стороны Заказчика.'
    )
    doc.add_paragraph(
        '3.2. Изготовление и поставка продукции осуществляется в соответствии с переданным и утвержденным Заказчиком оригинал-макетом.'
    )
    doc.add_paragraph(
        '3.3. При отсутствии у Заказчика оригинал-макета в готовом виде, Исполнитель по поручению Заказчика принимает на себя обязательство '
        'по изготовлению оригинал-макета за дополнительную плату. Заказчик должен осмотреть и подписать (либо указать на требуемые исправления) '
        'разработанный и предоставленный оригинал-макет в течение 7 (семи) рабочих дней с момента предоставления.'
    )
    doc.add_paragraph(
        '3.4. Исполнитель приступает к изготовлению продукции и в установленные сроки извещает Заказчика о готовности продукции к передаче.'
    )
    doc.add_paragraph(
        '3.5. Предварительная проверка продукции по качеству и количеству осуществляется уполномоченными представителями Заказчика на складе Исполнителя. '
        'Представитель Заказчика при получении продукции должен предъявить паспорт и надлежаще оформленную доверенность.'
    )
    doc.add_paragraph(
        '3.6. Продукция считается переданной Исполнителем и принятой Заказчиком в момент фактической передачи продукции и подписания накладной и счет-фактуры.'
    )

    # Заголовок раздела 4
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('4. ПРАВА И ОБЯЗАННОСТИ СТОРОН')
    run.bold = True

    doc.add_paragraph('4.1. Исполнитель обязан:')
    doc.add_paragraph('4.1.1. Выполнить принятые им обязательства в соответствии с настоящим Договором.')
    doc.add_paragraph('4.1.2. Передать уполномоченному лицу Заказчика продукцию в полном объеме.')

    doc.add_paragraph('4.2. Исполнитель вправе:')
    doc.add_paragraph(
        '4.2.1. Привлекать третьих лиц для выполнения принятых обязательств, оставаясь ответственным перед Заказчиком.'
    )
    doc.add_paragraph(
        '4.2.2. Запрашивать у Заказчика информацию, необходимую для выполнения работ.'
    )
    doc.add_paragraph(
        '4.2.3. Требовать от Заказчика надлежащего исполнения обязательств по Договору.'
    )

    doc.add_paragraph('4.3. Заказчик обязан:')
    doc.add_paragraph('4.3.1. Оплатить работы в соответствии с пунктами 2.1 и 2.2.')
    doc.add_paragraph('4.3.2. Принять изготовленную Исполнителем продукцию в соответствии с настоящим Договором.')

    doc.add_paragraph('4.4. Заказчик вправе:')
    doc.add_paragraph(
        '4.4.1. В любое время проверять ход и качество выполнения работы, не вмешиваясь в деятельность Исполнителя.'
    )

    # Заголовок раздела 5
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('5. ОСОБЫЕ УСЛОВИЯ')
    run.bold = True

    doc.add_paragraph(
        '5.1. Исполнитель даёт гарантию на изготовление качественной продукции и срок её годности в течение 12 месяцев.'
    )
    doc.add_paragraph(
        '5.2. В случае аннулирования заказа после начала изготовления, Заказчик обязан пропорционально оплатить все понесённые расходы.'
    )
    doc.add_paragraph(
        '5.3. Претензии по качеству продукции могут быть предъявлены не позднее 5 дней с момента её получения.'
    )
    doc.add_paragraph(
        '5.4. При наличии обоснованных претензий по качеству, Исполнитель обязуется в 15-дневный срок заменить продукцию за свой счёт.'
    )
    doc.add_paragraph(
        '5.5.  В случае несогласия ЗАКАЗЧИКА с изменением Цены продукции или прочие условия договора, СТОРОНЫ вправе отказаться от исполнения обязательств по настоящему Договору и расторгнуть его в одностороннем порядке. Договор будет считаться расторгнутым'
    )

    doc.add_paragraph(
        'по истечении 5 (пяти) рабочих дней с момента получения ЗАКАЗЧИКОМ  письменного уведомления ИСПОЛНИТЕЛЮ о его расторжении'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('6. ОТВЕТСТВЕННОСТЬ СТОРОН')
    run.bold = True

    doc.add_paragraph(
        '6.1. За не исполнение и(или) нарушение сроков оплаты, предусмотренных пункта 2.2. Заказчик уплачивает Исполнителю пеню в размере 0,4% от суммы просроченного платежа за каждый день просрочки, но не более  50% от суммы просроченного платежа.'
    )

    doc.add_paragraph(
        '6.2. За не исполнение и(или) нарушение сроков поставки продукции, согласно пункта 3.1. Исполнитель уплачивает Заказчику пеню в размере 0,5% от суммы просроченного платежа за каждый день просрочки, но не более 50% от суммы просроченного платежа.'
    )

    # Заголовок раздела 7
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('7. ДЕЙСТВИЕ ОБСТОЯТЕЛЬСТВ НЕПРЕОДОЛИМОЙ СИЛЫ (ФОРС-МАЖОР)')
    run.bold = True


    doc.add_paragraph(
        '7.1. СТОРОНЫ освобождаются от ответственности за частичное или полное неисполнение обязательств по настоящему Договору, если это неисполнение явилось следствием обстоятельств непреодолимой силы, возникших после заключения настоящего Договора в результате событий чрезвычайного характера, которые СТОРОНА не могла ни предвидеть, ни предотвратить разумными мерами (форс-мажор). К таким событиям чрезвычайного характера относятся: наводнение, пожар, землетрясение, взрыв, шторм, оседание почвы, эпидемии и иные явления природы, а также война или военные действия, забастовка в отрасли или регионе, принятие органом государственной власти или управления решения, повлекшего невозможность исполнения настоящего договора.'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('8. ПРОЧИЕ УСЛОВИЯ')
    run.bold = True

    doc.add_paragraph(
        '8.1. В случае возникновения разногласий все споры решаются путем двухсторонних переговоров, а при невозможности прийти к соглашению, все споры подлежат рассмотрению экономическим судом города Ташкента.'
    )

    doc.add_paragraph(
        '8.2. Все изменения и дополнения к настоящему Договору действительны, если они составлены в письменной форме и подписаны Сторонами.'
    )

    doc.add_paragraph(
        '8.3. Настоящий договор составлен в двух экземплярах, по одному для каждой из сторон, имеющих одинаковую юридическую силу.'
    )

    doc.add_paragraph(
        '8.4.При возникновении споров и разногласий  Стороны руководствуются действующим законодательством РУз.'
    )

    # Заголовок раздела 9
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('9. СРОК ДЕЙСТВИЯ ДОГОВОРА')
    run.bold = True


    doc.add_paragraph(
        '9.1.Настоящий договор вступает в силу со дня  регистрации   и действует до  31.12.2025г.'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('10. АДРЕСА, РЕКВИЗИТЫ И ПОДПИСИ СТОРОН')
    run.bold = True
    
    # Создание таблицы
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.2)
    table.columns[1].width = Inches(3.2)

    # Ячейка 1 — ИСПОЛНИТЕЛЬ
    cell1 = table.cell(0, 0)
    cell1.text = ''  # Очищаем содержимое
    p = cell1.paragraphs[0]
    run = p.add_run(
        '\tИСПОЛНИТЕЛЬ:\n'
        'ООО «NUR FAYZ REKLAMA»\n'
        'Адрес: г. Ташкент, Юкоричирчикский район,\n'
        'Ийк-ота ота МФЙ, р/з- уй\n'
        'р/с: 2020 8000 2055 4211 2001\n'
        'в г. АНДИЖАН., ГОЛОВНОЙ ОФИС АК-БАНКА «HAMKORBANK»\n'
        'С УЧАСТИЕМ ИНОСТРАННОГО КАПИТАЛА\n'
        'МФО: 00083   ИНН: 309667279\n'
        'Телефон: (99899) 804-30-02\n\n'
        'Директор _______________ Файзиев Н.С.\n'
    )

    # Добавляем изображение после текста
    p_img = cell1.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.LEFT  # или CENTER
    p_img.add_run().add_picture('1.png', width=Inches(2))  # Измени ширину при необходимости

    # Ячейка 2 — ЗАКАЗЧИК
    cell2 = table.cell(0, 1)
    cell2.text = ''
    p = cell2.paragraphs[0]
    p.add_run(
        '\tЗАКАЗЧИК:\n'
        'ООО «HIGH CITY-DEVELOPERS»\n'
        'Адрес: Тошкент вилояти, Зангиота тумани, Эркин МФЙ,\n'
        'Катта узбек тракти, 5-уй\n'
        'р/с: 2020 8000 8007 9248 2001\n'
        'Банк: ЧАКБ «DAVR-BANK» Юнусабадский ф-л\n'
        'МФО: 01122   ИНН: 304998124\n'
        'Код НДС: 303010207864\n\n'
        'Директор: _______________ Ишмурадов Ш.Қ.'
    )

    # Вставляем картинку на полный лист
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture('2.png', width=Cm(15.21), height=Cm(22))



    # Вставляем картинку на полный лист
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture('3.png', width=Cm(15.21), height=Cm(22))

    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    return excel_buffer, docx_buffer
